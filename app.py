from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'

# Шаблоны данных для документов
DOCUMENT_TYPES = {
    'policy': 'Политика обработки персональных данных',
    'cookie': 'Политика использования cookie',
    'consent_processing': 'Согласие на обработку персональных данных',
    'consent_newsletter': 'Согласие на рассылку',
    'consent_distribution': 'Согласие на распространение ПДн',
    'user_agreement': 'Пользовательское соглашение'
}

SITE_FEATURES = {
    'forms': 'Формы сбора данных (контактные, заявки)',
    'analytics': 'Системы аналитики (Яндекс.Метрика, Google Analytics)',
    'cookies': 'Использование cookie-файлов',
    'newsletter': 'Email-рассылка',
    'third_party': 'Передача данных третьим лицам',
    'crm': 'Интеграция с CRM-системами',
    'payment': 'Платежные системы',
    'social': 'Виджеты социальных сетей'
}

def generate_policy_doc(data):
    """Генерация политики обработки персональных данных"""
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('ПОЛИТИКА ОБРАБОТКИ ПЕРСОНАЛЬНЫХ ДАННЫХ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'г. {data.get("city", "______")}')
    doc.add_paragraph(f'«{data.get("date", "______")}» {data.get("year", "2024")} г.')
    
    # Разделы
    doc.add_heading('1. ОБЩИЕ ПОЛОЖЕНИЯ', level=1)
    doc.add_paragraph(
        f'Настоящая Политика обработки персональных данных (далее — Политика) действует в отношении '
        f'всех персональных данных, которые {data.get("org_name", "Организация")} (далее — Оператор) '
        f'может получить от посетителей веб-сайта {data.get("site_url", "www.example.ru")}.'
    )
    
    doc.add_heading('2. ЦЕЛИ ОБРАБОТКИ ПЕРСОНАЛЬНЫХ ДАННЫХ', level=1)
    purposes = []
    if 'forms' in data.get('features', []):
        purposes.append('- обработка заявок и обращений пользователей')
    if 'analytics' in data.get('features', []):
        purposes.append('- анализ посещаемости сайта и улучшение качества обслуживания')
    if 'newsletter' in data.get('features', []):
        purposes.append('- направление информационных сообщений и рекламных материалов')
    if 'payment' in data.get('features', []):
        purposes.append('- обработка платежей и выполнение финансовых операций')
    
    for purpose in purposes:
        doc.add_paragraph(purpose)
    
    doc.add_heading('3. ПРАВОВЫЕ ОСНОВАНИЯ ОБРАБОТКИ', level=1)
    doc.add_paragraph(
        'Обработка персональных данных осуществляется на основании:\n'
        '- Федерального закона от 27.07.2006 № 152-ФЗ «О персональных данных»;\n'
        '- согласия субъекта персональных данных;\n'
        '- договора с субъектом персональных данных.'
    )
    
    doc.add_heading('4. ОБЪЕМ И КАТЕГОРИИ ПЕРСОНАЛЬНЫХ ДАННЫХ', level=1)
    doc.add_paragraph(
        f'Оператор может обрабатывать следующие персональные данные:\n'
        f'- ФИО: {data.get("collect_fio", "нет")}\n'
        f'- Email: {data.get("collect_email", "нет")}\n'
        f'- Телефон: {data.get("collect_phone", "нет")}\n'
        f'- Адрес: {data.get("collect_address", "нет")}'
    )
    
    doc.add_heading('5. МЕРЫ ЗАЩИТЫ', level=1)
    doc.add_paragraph(
        'Оператор принимает необходимые организационные и технические меры для защиты '
        'персональных данных от неправомерного доступа, уничтожения, изменения, блокирования, '
        'копирования, предоставления, распространения.'
    )
    
    doc.add_heading('6. КОНТАКТНАЯ ИНФОРМАЦИЯ', level=1)
    doc.add_paragraph(
        f'Адрес: {data.get("org_address", "______")}\n'
        f'Email: {data.get("org_email", "______")}\n'
        f'Телефон: {data.get("org_phone", "______")}'
    )
    
    return doc

def generate_cookie_doc(data):
    """Генерация политики использования cookie"""
    doc = Document()
    
    title = doc.add_heading('ПОЛИТИКА ИСПОЛЬЗОВАНИЯ COOKIE-ФАЙЛОВ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Веб-сайт: {data.get("site_url", "www.example.ru")}')
    doc.add_paragraph(f'Дата обновления: {data.get("date", "______")} {data.get("year", "2024")} г.')
    
    doc.add_heading('1. ЧТО ТАКОЕ COOKIE-ФАЙЛЫ', level=1)
    doc.add_paragraph(
        'Cookie-файлы — это небольшие текстовые файлы, которые сохраняются на устройстве пользователя '
        'при посещении веб-сайта. Они содержат информацию о действиях пользователя на сайте.'
    )
    
    doc.add_heading('2. КАКИЕ COOKIE МЫ ИСПОЛЬЗУЕМ', level=1)
    
    if 'analytics' in data.get('features', []):
        doc.add_heading('2.1. Аналитические cookie', level=2)
        doc.add_paragraph(
            'Мы используем Яндекс.Метрику и/или Google Analytics для анализа посещаемости сайта. '
            'Эти сервисы устанавливают cookie-файлы для сбора статистики.'
        )
    
    if 'forms' in data.get('features', []):
        doc.add_heading('2.2. Функциональные cookie', level=2)
        doc.add_paragraph(
            'Необходимы для работы форм обратной связи и сохранения настроек пользователя.'
        )
    
    doc.add_heading('3. УПРАВЛЕНИЕ COOKIE', level=1)
    doc.add_paragraph(
        'Вы можете управлять настройками cookie в своем браузере. Большинство браузеров позволяют '
        'блокировать или удалять cookie-файлы.'
    )
    
    doc.add_heading('4. КОНТАКТЫ', level=1)
    doc.add_paragraph(f'Email для вопросов: {data.get("org_email", "______")}')
    
    return doc

def generate_consent_processing_doc(data):
    """Генерация согласия на обработку персональных данных"""
    doc = Document()
    
    title = doc.add_heading('СОГЛАСИЕ НА ОБРАБОТКУ ПЕРСОНАЛЬНЫХ ДАННЫХ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        f'Я, нижеподписавшийся(аяся), даю свое согласие {data.get("org_name", "Организация")} '
        f'(ИНН {data.get("inn", "______")}), расположенной по адресу: {data.get("org_address", "______")}, '
        f'на обработку моих персональных данных.'
    )
    
    doc.add_heading('1. ПЕРЕЧЕНЬ ДАННЫХ', level=1)
    items = []
    if data.get("collect_fio") == "да":
        items.append('- ФИО')
    if data.get("collect_email") == "да":
        items.append('- Адрес электронной почты')
    if data.get("collect_phone") == "да":
        items.append('- Номер телефона')
    if data.get("collect_address") == "да":
        items.append('- Почтовый адрес')
    
    for item in items:
        doc.add_paragraph(item)
    
    doc.add_heading('2. ЦЕЛИ ОБРАБОТКИ', level=1)
    doc.add_paragraph(
        'Обработка данных осуществляется для:\n'
        '- заключения и исполнения договоров;\n'
        '- направления информационных сообщений;\n'
        '- улучшения качества обслуживания.'
    )
    
    doc.add_heading('3. СПОСОБЫ ОБРАБОТКИ', level=1)
    doc.add_paragraph(
        'Обработка может осуществляться как автоматизированным, так и неавтоматизированным способом.'
    )
    
    doc.add_paragraph('\n\nСогласие дано «______» _____________ 20___ г.')
    doc.add_paragraph('\nПодпись: _________________')
    
    return doc

def generate_consent_newsletter_doc(data):
    """Генерация согласия на рассылку"""
    doc = Document()
    
    title = doc.add_heading('СОГЛАСИЕ НА ПОЛУЧЕНИЕ РАССЫЛКИ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        f'Я выражаю согласие {data.get("org_name", "Организация")} на получение информационных '
        f'и рекламных сообщений на адрес электронной почты: {data.get("user_email", "______")}.'
    )
    
    doc.add_paragraph(
        'Я понимаю, что могу отказаться от получения рассылки в любой момент, '
        'нажав на ссылку «Отписаться» в письме или направив запрос на email: '
        f'{data.get("org_email", "______")}.'
    )
    
    doc.add_paragraph('\nДата: «______» _____________ 20___ г.')
    doc.add_paragraph('Подпись: _________________')
    
    return doc

def generate_consent_distribution_doc(data):
    """Генерация согласия на распространение ПДн"""
    doc = Document()
    
    title = doc.add_heading('СОГЛАСИЕ НА РАСПРОСТРАНЕНИЕ ПЕРСОНАЛЬНЫХ ДАННЫХ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        f'Я, нижеподписавшийся(аяся), даю согласие {data.get("org_name", "Организация")} '
        f'на распространение моих персональных данных в соответствии со ст. 10.1 ФЗ-152.'
    )
    
    doc.add_heading('Перечень данных, разрешенных к распространению:', level=1)
    items = []
    if data.get("collect_fio") == "да":
        items.append('- ФИО')
    if data.get("collect_email") == "да":
        items.append('- Email')
    if data.get("collect_phone") == "да":
        items.append('- Телефон')
    
    for item in items:
        doc.add_paragraph(item)
    
    doc.add_paragraph(
        '\nЯ подтверждаю, что понимание содержания настоящего согласия у меня имеется.'
    )
    
    doc.add_paragraph('\nДата: «______» _____________ 20___ г.')
    doc.add_paragraph('Подпись: _________________')
    
    return doc

def generate_user_agreement_doc(data):
    """Генерация пользовательского соглашения"""
    doc = Document()
    
    title = doc.add_heading('ПОЛЬЗОВАТЕЛЬСКОЕ СОГЛАШЕНИЕ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Веб-сайт: {data.get("site_url", "www.example.ru")}')
    doc.add_paragraph(f'Дата публикации: {data.get("date", "______")} {data.get("year", "2024")} г.')
    
    doc.add_heading('1. ОБЩИЕ ПОЛОЖЕНИЯ', level=1)
    doc.add_paragraph(
        f'Настоящее Пользовательское соглашение (далее — Соглашение) регулирует отношения между '
        f'{data.get("org_name", "Организацией")} и пользователем веб-сайта {data.get("site_url", "www.example.ru")}.'
    )
    
    doc.add_heading('2. ПРЕДМЕТ СОГЛАШЕНИЯ', level=1)
    doc.add_paragraph(
        '2.1. Владелец сайта предоставляет пользователю доступ к информации и сервисам сайта.\n'
        '2.2. Использование сайта означает полное и безоговорочное принятие условий Соглашения.'
    )
    
    doc.add_heading('3. ОБЯЗАТЕЛЬСТВА СТОРОН', level=1)
    doc.add_paragraph(
        '3.1. Владелец обязуется предоставлять достоверную информацию.\n'
        '3.2. Пользователь обязуется не нарушать работу сайта и использовать его в законных целях.'
    )
    
    if 'payment' in data.get('features', []):
        doc.add_heading('4. ОПЛАТА И ВОЗВРАТ', level=1)
        doc.add_paragraph(
            '4.1. Оплата товаров/услуг производится в соответствии с выбранным способом оплаты.\n'
            '4.2. Возврат средств осуществляется в порядке, установленном законодательством РФ.'
        )
    
    doc.add_heading('5. ОТВЕТСТВЕННОСТЬ', level=1)
    doc.add_paragraph(
        '5.1. Владелец не несет ответственности за убытки, возникшие вследствие использования сайта.\n'
        '5.2. Владелец вправе изменять условия Соглашения в одностороннем порядке.'
    )
    
    doc.add_heading('6. РЕКВИЗИТЫ ВЛАДЕЛЬЦА', level=1)
    doc.add_paragraph(
        f'Наименование: {data.get("org_name", "______")}\n'
        f'ИНН: {data.get("inn", "______")}\n'
        f'Адрес: {data.get("org_address", "______")}\n'
        f'Email: {data.get("org_email", "______")}\n'
        f'Телефон: {data.get("org_phone", "______")}'
    )
    
    return doc

GENERATORS = {
    'policy': generate_policy_doc,
    'cookie': generate_cookie_doc,
    'consent_processing': generate_consent_processing_doc,
    'consent_newsletter': generate_consent_newsletter_doc,
    'consent_distribution': generate_consent_distribution_doc,
    'user_agreement': generate_user_agreement_doc
}

@app.route('/')
def index():
    return render_template('index.html', 
                         document_types=DOCUMENT_TYPES, 
                         site_features=SITE_FEATURES)

@app.route('/generate', methods=['POST'])
def generate():
    try:
        doc_type = request.form.get('doc_type')
        if not doc_type or doc_type not in GENERATORS:
            return 'Неверный тип документа', 400
        
        # Сбор данных из формы
        data = {
            'doc_type': doc_type,
            'org_name': request.form.get('org_name', ''),
            'inn': request.form.get('inn', ''),
            'org_address': request.form.get('org_address', ''),
            'org_email': request.form.get('org_email', ''),
            'org_phone': request.form.get('org_phone', ''),
            'site_url': request.form.get('site_url', ''),
            'city': request.form.get('city', ''),
            'date': request.form.get('date', ''),
            'year': request.form.get('year', '2024'),
            'features': request.form.getlist('features'),
            'collect_fio': request.form.get('collect_fio', 'нет'),
            'collect_email': request.form.get('collect_email', 'нет'),
            'collect_phone': request.form.get('collect_phone', 'нет'),
            'collect_address': request.form.get('collect_address', 'нет'),
            'user_email': request.form.get('user_email', '')
        }
        
        # Генерация документа
        generator = GENERATORS[doc_type]
        doc = generator(data)
        
        # Сохранение в буфер
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"{doc_type}_{data['org_name'].replace(' ', '_')}.docx"
        
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    
    except Exception as e:
        return f'Ошибка генерации: {str(e)}', 500

if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    app.run(debug=False, host='0.0.0.0', port=5000, use_reloader=False)
